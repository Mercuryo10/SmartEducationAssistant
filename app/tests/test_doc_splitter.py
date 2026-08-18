"""多格式分块器单测（阶段 7.1，docs/08 §10）。

覆盖 md / txt / html / docx 四种格式的硬性约束：
1. 一个文本块只属于一个标题，不跨标题。
2. 块文本前拼接完整标题路径（# 一级 \\n ## 二级）。
3. 长段落经 RecursiveCharacterTextSplitter 二次切分后，各子块共享同一标题路径。
4. 不支持格式（如 pdf）抛结构化 ValidationError。

纯文本离线运行：docx 用 python-docx 在 tmp_path 现场构造，不依赖真实文件。
"""
import tempfile
from pathlib import Path

import pytest

from app.core.exceptions import ValidationError
from app.services.doc_splitter import (
    extract_title,
    extract_title_file,
    parse_docx_sections,
    parse_html_sections,
    parse_md_sections,
    split_document,
    split_document_file,
    split_document_file_sections,
    split_document_sections,
)


def _make_docx(path: Path, items: list[tuple[int, str]]) -> None:
    """构造 docx：items 为 (标题级别 0=普通段落, 文本)。"""
    from docx import Document

    doc = Document()
    for level, text in items:
        if level:
            doc.add_heading(text, level=level)
        else:
            doc.add_paragraph(text)
    doc.save(path)


# ---------- markdown（md 格式，原 md_splitter 用例迁移） ----------

def test_one_chunk_per_section() -> None:
    """每节一个块：块与 ## 小节一一对应，不跨标题。"""
    md = "# 文档\n\n## 甲\n\n内容甲。\n\n## 乙\n\n内容乙。\n\n## 丙\n\n内容丙。"
    chunks = split_document_sections(md, "md")
    assert [c["heading_path"] for c in chunks] == [
        "# 文档\n## 甲",
        "# 文档\n## 乙",
        "# 文档\n## 丙",
    ]


def test_heading_path_prepended() -> None:
    """块文本以完整标题路径开头（# 一级 \\n ## 二级 \\n \\n 正文）。"""
    md = "# 检索增强生成（RAG）原理\n\n## 1. 什么是 RAG\n\nRAG 解决的核心问题是幻觉。"
    chunks = split_document_sections(md, "md")
    assert chunks[0]["text"].startswith("# 检索增强生成（RAG）原理\n## 1. 什么是 RAG\n\n")
    assert "幻觉" in chunks[0]["text"]


def test_long_section_split_keeps_heading_path() -> None:
    """长段落被二次切分，但所有子块仍只属同一标题且路径相同。"""
    md = "# 总标题\n\n## 长小节\n\n" + "这是第一句。" * 40  # 约 240 字
    chunks = split_document_sections(md, "md", chunk_size=100, chunk_overlap=10)
    assert len(chunks) >= 2, "超过 chunk_size 的段落应被切分"
    for c in chunks:
        assert c["heading_path"] == "# 总标题\n## 长小节"
        assert c["text"].startswith("# 总标题\n## 长小节\n\n")


def test_short_section_single_chunk() -> None:
    """未超 chunk_size 的小节保持一块。"""
    md = "# 标题\n\n## 小节\n\n短短一段正文。"
    chunks = split_document_sections(md, "md")
    assert len(chunks) == 1
    assert chunks[0]["heading_path"] == "# 标题\n## 小节"


def test_intro_before_first_section_belongs_to_title() -> None:
    """首个 # 标题下的引导正文归属该标题，路径不含 ##。"""
    md = "# 总览\n\n这是总览正文。\n\n## 分节\n\n分节正文。"
    chunks = split_document_sections(md, "md")
    assert chunks[0]["heading_path"] == "# 总览"
    assert "总览正文" in chunks[0]["text"]


def test_three_level_heading_path() -> None:
    """三级标题路径逐级拼接。"""
    md = "# 一级\n\n## 二级\n\n### 三级\n\n正文。"
    chunks = split_document_sections(md, "md")
    assert chunks[0]["heading_path"] == "# 一级\n## 二级\n### 三级"


def test_no_heading_text_kept_as_chunk() -> None:
    """全文无标题的文本不丢数据：产出 heading_path 为空的单块。"""
    assert split_document("没有标题的纯文本。", "txt") == ["没有标题的纯文本。"]


def test_parse_sections_skips_empty_section() -> None:
    """空小节（标题下无正文）不产生分块。"""
    md = "# 文档\n\n## 空小节\n\n## 非空\n\n有内容。"
    chunks = split_document_sections(md, "md")
    assert len(chunks) == 1
    assert chunks[0]["heading_path"] == "# 文档\n## 非空"


def test_split_document_returns_text_list() -> None:
    """split_document 是结构化分块的 text 扁平化。"""
    md = "# 标题\n\n## 小节\n\n正文。"
    assert split_document(md, "md") == ["# 标题\n## 小节\n\n正文。"]


def test_parse_md_sections_structure() -> None:
    """parse_md_sections 返回节的标题路径与正文。"""
    sections = parse_md_sections("# A\n\n## B\n\n内容")
    assert sections[0]["heading_path"] == "# A\n## B"
    assert sections[0]["content"].strip() == "内容"


# ---------- txt ----------

def test_txt_with_hash_headings_is_heading_aware() -> None:
    """txt 里用 `#` 标题行即可标题感知（与 md 同逻辑）。"""
    txt = "# Python 学习笔记\n\n## 1. 基础语法\n\nPython 是解释型语言。"
    chunks = split_document_sections(txt, "txt")
    assert chunks[0]["heading_path"] == "# Python 学习笔记\n## 1. 基础语法"


# ---------- html ----------

def test_html_sections_from_h_tags() -> None:
    """h1-h2 层级解析为节列表，正文按块归属。"""
    html = (
        "<html><body><h1>HTML 知识图谱</h1><p>导语。</p>"
        "<h2>1. 实体</h2><p>实体是节点。</p>"
        "<h2>2. 关系</h2><p>关系是边。</p></body></html>"
    )
    chunks = split_document_sections(html, "html")
    assert [c["heading_path"] for c in chunks] == [
        "# HTML 知识图谱",
        "# HTML 知识图谱\n## 1. 实体",
        "# HTML 知识图谱\n## 2. 关系",
    ]
    assert chunks[0]["text"].startswith("# HTML 知识图谱\n\n")


def test_html_script_style_noise_removed() -> None:
    """script/style 内容不进入知识库分块。"""
    html = "<h1>标题</h1><p>正文。</p><script>var secret = 1;</script><style>.x{}</style>"
    chunks = split_document_sections(html, "html")
    assert all("secret" not in c["text"] and ".x{}" not in c["text"] for c in chunks)


def test_parse_html_sections_returns_sections() -> None:
    sections = parse_html_sections("<h1>甲</h1><p>正文</p>")
    assert sections[0]["heading_path"] == "# 甲"
    assert sections[0]["content"].strip() == "正文"


def test_html_doctype_and_comment_ignored() -> None:
    """<!DOCTYPE html> 与注释不进入知识库正文（回归：Doctype 的字符串是 'html'）。"""
    html = "<!DOCTYPE html><html><body><h1>标题</h1><!-- 注释 --><p>正文。</p></body></html>"
    chunks = split_document_sections(html, "html")
    assert [c["heading_path"] for c in chunks] == ["# 标题"]
    assert all("DOCTYPE" not in c["text"] and "注释" not in c["text"] for c in chunks)


# ---------- docx ----------

def test_docx_sections_from_heading_styles() -> None:
    """Word 内置标题样式映射为节，普通段落归属到当前节。"""
    p = Path(tempfile.mkdtemp()) / "sample.docx"
    _make_docx(
        p,
        [
            (1, "Word 教材示例"),
            (0, "这是正文。"),
            (2, "第一节：标题样式"),
            (0, "段落内容甲。"),
            (2, "第二节：分块验证"),
            (0, "段落内容乙。"),
        ],
    )
    sections = parse_docx_sections(p)
    assert [s["heading_path"] for s in sections] == [
        "# Word 教材示例",
        "# Word 教材示例\n## 第一节：标题样式",
        "# Word 教材示例\n## 第二节：分块验证",
    ]
    assert "这是正文" in sections[0]["content"]
    assert "段落内容乙" in sections[2]["content"]


def test_docx_chunks_prepend_heading_path() -> None:
    p = Path(tempfile.mkdtemp()) / "sample.docx"
    _make_docx(p, [(1, "教材"), (0, "正文。"), (2, "节"), (0, "节内容。")])
    chunks = split_document_file_sections(p)
    assert len(chunks) == 2
    assert chunks[1]["text"].startswith("# 教材\n## 节\n\n")


def test_docx_without_styles_single_section() -> None:
    """docx 未用标题样式（仅普通段落）时整篇一节，不丢数据。"""
    p = Path(tempfile.mkdtemp()) / "plain.docx"
    _make_docx(p, [(0, "没有标题样式的正文一。"), (0, "正文二。")])
    sections = parse_docx_sections(p)
    assert len(sections) == 1
    assert sections[0]["heading_path"] == ""
    assert "正文一" in sections[0]["content"]


# ---------- 分发与错误路径 ----------

def test_split_document_unsupported_format_raises() -> None:
    with pytest.raises(ValidationError):
        split_document("内容", "pdf")


def test_split_document_file_unsupported_suffix_raises() -> None:
    with pytest.raises(ValidationError):
        split_document_file("教材.pdf")


def test_split_document_file_txt() -> None:
    p = Path(tempfile.mkdtemp()) / "notes.txt"
    p.write_text("# 标题\n\n## 小节\n\n正文。", encoding="utf-8")
    chunks = split_document_file(p)
    assert chunks == ["# 标题\n## 小节\n\n正文。"]


# ---------- 标题提取 ----------

def test_extract_title_md_and_html() -> None:
    assert extract_title("# 我的标题\n\n正文", "md") == "我的标题"
    assert extract_title("<h1>HTML 标题</h1><p>正文</p>", "html") == "HTML 标题"
    # 无 h1 时回退 <title>
    assert extract_title("<title>仅标题页</title><p>正文</p>", "html") == "仅标题页"
    # 无标题的 txt 回退首个非空行
    assert extract_title("首行内容\n第二行", "txt") == "首行内容"


def test_extract_title_file_docx() -> None:
    p = Path(tempfile.mkdtemp()) / "title.docx"
    _make_docx(p, [(1, "Doc 标题"), (0, "正文。")])
    assert extract_title_file(p) == "Doc 标题"
