"""多格式文本分块器（阶段 7.1 扩展，docs/09 待补）。

把 md / txt / html / docx 统一为**标题感知分块**：
1. 每种格式先解析为节列表 [{heading_path, level, content}]（heading_path = 完整标题链）。
2. 节内正文按标题边界用 RecursiveCharacterTextSplitter 二次切分（含中文标点断点）。
3. 每块文本前拼接所属标题路径，检索命中时自带上下文归属。

各格式的标题信号：
- md / txt：`#` 标题行；txt 无标题时整篇一节（不丢数据）。
- html：h1-h6 标签层级（自动剔除 script/style/head 噪声）。
- docx：Word 内置「标题 N」样式；仅手动加粗/放大字号视为普通段落，无标题信号。
- pdf：暂不支持（文档格式解析留待后续，此处直接抛结构化错误）。

bs4 / python-docx 均为惰性导入（仅对应解析器用到），不装不影响其余格式与启动。
"""
import re
from pathlib import Path
from typing import Any

from langchain_text_splitters import RecursiveCharacterTextSplitter

from app.core.config import settings
from app.core.exceptions import ValidationError

# 标题行：1-6 个 # 后跟空格与标题文本（md/txt 约定，html/docx 转换后同样式）
_HEADING_RE = re.compile(r"^(#{1,6})\s+(.+?)\s*$")

# 二次切分分隔符：段落 / 行 / 中文句读 / 空格（langchain 按顺序优先匹配）
_MD_SEPARATORS = ["\n\n", "\n", "。", "；", "，", " ", ""]

_HTML_HEADING_TAGS = {"h1", "h2", "h3", "h4", "h5", "h6"}
_HTML_BLOCK_TAGS = {"p", "div", "li", "ul", "ol", "tr", "section", "article", "blockquote", "br", "pre", "table"}
_DOCX_HEADING_RE = re.compile(r"Heading (\d+)")

# 支持的文件后缀 → 格式名（pdf 不在其中，视为不支持）
_SUFFIX_TO_FORMAT = {
    ".md": "md",
    ".markdown": "md",
    ".txt": "txt",
    ".html": "html",
    ".htm": "html",
    ".docx": "docx",
}


# ---------- 各格式解析为节列表 ----------

def parse_md_sections(text: str) -> list[dict[str, Any]]:
    """把 markdown / 带 `#` 标题的 txt 解析为节列表。

    Args:
        text: 文档全文。

    Returns:
        每节 {heading_path, level, content}；首个标题前的正文并入其下第一个节；
        全文无标题时 heading_path 为空串（不丢数据）。
    """
    sections: list[dict[str, Any]] = []
    stack: list[str] = []  # 标题路径栈（含 # 前缀，逐级截断）
    cur: dict[str, Any] = {"heading_path": "", "level": 0, "content": ""}
    for line in text.splitlines():
        m = _HEADING_RE.match(line)
        if m:
            if cur["content"].strip():
                sections.append(cur)
            level = len(m.group(1))
            stack = stack[: level - 1]  # 截断到父级，保持路径正确
            stack.append(line.strip())
            cur = {"heading_path": "\n".join(stack), "level": level, "content": ""}
        else:
            cur["content"] += line + "\n"
    if cur["content"].strip():
        sections.append(cur)
    return sections


def parse_txt_sections(text: str) -> list[dict[str, Any]]:
    """txt 解析：复用 md 标题识别；无 `#` 标题时整篇一节。"""
    return parse_md_sections(text)


def parse_html_sections(text: str) -> list[dict[str, Any]]:
    """解析 HTML 的 h1-h6 层级为节列表。

    Args:
        text: html 文档文本。

    Returns:
        每节 {heading_path, level, content}；无标题时 heading_path 为空串。
        剔除 script/style/head 噪声（正文外的脚本、元信息不进入知识库）。
    """
    from bs4 import BeautifulSoup, Comment, Doctype  # 惰性导入

    soup = BeautifulSoup(text, "html.parser")
    for tag in soup(["script", "style", "head"]):
        tag.decompose()

    sections: list[dict[str, Any]] = []
    stack: list[str] = []

    def walk(node: Any, cur: dict[str, Any]) -> dict[str, Any]:
        """递归遍历：标题截断栈并开新节，块级元素换行，文本节点归入当前节。"""
        if isinstance(node, (Doctype, Comment)):
            return cur  # 文档类型声明/注释无知识价值，跳过
        name = getattr(node, "name", None)
        if name in _HTML_HEADING_TAGS:
            if cur["content"].strip():
                sections.append(cur)
            level = int(name[1])
            new_stack = stack[: level - 1]  # 截断到父级，保持路径正确
            new_stack.append(f"{'#' * level} {node.get_text(' ', strip=True)}")
            stack.clear()
            stack.extend(new_stack)
            return {"heading_path": "\n".join(stack), "level": level, "content": ""}
        if name in _HTML_BLOCK_TAGS:
            cur["content"] += "\n"  # 块级元素换行，避免相邻文本粘连
        if isinstance(node, str):
            cur["content"] += node
            return cur
        for child in node.children:
            cur = walk(child, cur)
        return cur

    cur: dict[str, Any] = {"heading_path": "", "level": 0, "content": ""}
    for child in soup.children:
        cur = walk(child, cur)
    if cur["content"].strip():
        sections.append(cur)
    return sections


def _open_docx(path: str | Path):
    """惰性加载并打开 .docx，失败转结构化异常。"""
    from docx import Document  # 惰性导入

    try:
        return Document(str(path))
    except Exception as exc:
        raise ValidationError(f"docx 解析失败 {path}: {exc}") from exc


def parse_docx_sections(path: str | Path) -> list[dict[str, Any]]:
    """解析 Word 文档的内置「标题 N」样式为节列表。

    Args:
        path: .docx 文件路径。

    Returns:
        每节 {heading_path, level, content}；未用标题样式（仅加粗/放大字号）时整篇一节。
    """
    doc = _open_docx(path)
    sections: list[dict[str, Any]] = []
    stack: list[str] = []
    cur: dict[str, Any] = {"heading_path": "", "level": 0, "content": ""}
    for para in doc.paragraphs:
        text = para.text.strip()
        style = para.style.name if para.style else ""
        m = _DOCX_HEADING_RE.fullmatch(style or "")
        if m:
            if cur["content"].strip():
                sections.append(cur)
            level = int(m.group(1))
            stack = stack[: level - 1]
            stack.append(f"{'#' * level} {text}")
            cur = {"heading_path": "\n".join(stack), "level": level, "content": ""}
        elif text:
            cur["content"] += text + "\n"
    if cur["content"].strip():
        sections.append(cur)
    return sections


# ---------- 标题边界分块（各格式共用） ----------

def split_sections(
    sections: list[dict[str, Any]], chunk_size: int | None = None, chunk_overlap: int | None = None
) -> list[dict[str, str]]:
    """把节列表按标题边界分块，每块 {heading_path, text}（text 已拼接标题路径）。

    Args:
        sections: 解析器产出的节列表。
        chunk_size: 每块字符数，默认 settings.chunk_size（500）。
        chunk_overlap: 相邻块重叠字符数，默认 settings.chunk_overlap（50）。

    Returns:
        分块列表；heading_path 相同即同属一个标题，块不跨标题。
    """
    size = chunk_size or settings.chunk_size
    ov = chunk_overlap or settings.chunk_overlap
    splitter = RecursiveCharacterTextSplitter(chunk_size=size, chunk_overlap=ov, separators=_MD_SEPARATORS)
    out: list[dict[str, str]] = []
    for sec in sections:
        body = sec["content"].strip()
        if not body:
            continue
        for piece in splitter.split_text(body):
            piece = piece.strip()
            if not piece:
                continue
            text_out = f"{sec['heading_path']}\n\n{piece}" if sec["heading_path"] else piece
            out.append({"heading_path": sec["heading_path"], "text": text_out})
    return out


# ---------- 按文本 / 文件分发 ----------

_TEXT_FORMATS = {"md": parse_md_sections, "txt": parse_txt_sections, "html": parse_html_sections}


def split_document_sections(
    text: str, fmt: str, chunk_size: int | None = None, chunk_overlap: int | None = None
) -> list[dict[str, str]]:
    """按文本格式解析并分块（md/txt/html；docx 用 split_document_file_sections）。

    Args:
        text: 文档文本。
        fmt: 格式名（md/txt/html）。
        chunk_size / chunk_overlap: 分块参数，默认取 settings。

    Returns:
        分块列表（每块 {heading_path, text}）。

    Raises:
        ValidationError: 不支持的格式。
    """
    parser = _TEXT_FORMATS.get((fmt or "").lower())
    if parser is None:
        raise ValidationError(f"不支持的分块格式：{fmt}（支持 md/txt/html/docx）")
    return split_sections(parser(text), chunk_size, chunk_overlap)


def split_document_file_sections(
    path: str | Path, chunk_size: int | None = None, chunk_overlap: int | None = None
) -> list[dict[str, str]]:
    """按文件后缀解析并分块（md/txt/html/docx，pdf 暂不支持）。

    Args:
        path: 文件路径。
        chunk_size / chunk_overlap: 分块参数，默认取 settings。

    Returns:
        分块列表（每块 {heading_path, text}）。

    Raises:
        ValidationError: 不支持的文件后缀（如 .pdf）。
    """
    p = Path(path)
    fmt = _SUFFIX_TO_FORMAT.get(p.suffix.lower())
    if fmt is None:
        raise ValidationError(f"不支持的知识库文件格式：{p.suffix}（支持 md/txt/html/docx）")
    if fmt == "docx":
        return split_sections(parse_docx_sections(p), chunk_size, chunk_overlap)
    return split_document_sections(p.read_text(encoding="utf-8"), fmt, chunk_size, chunk_overlap)


def split_document(text: str, fmt: str, chunk_size: int | None = None, chunk_overlap: int | None = None) -> list[str]:
    """按文本格式解析并分块，返回块文本列表（首行为所属标题路径）。"""
    return [d["text"] for d in split_document_sections(text, fmt, chunk_size, chunk_overlap)]


def split_document_file(
    path: str | Path, chunk_size: int | None = None, chunk_overlap: int | None = None
) -> list[str]:
    """按文件后缀解析并分块，返回块文本列表（首行为所属标题路径）。"""
    return [d["text"] for d in split_document_file_sections(path, chunk_size, chunk_overlap)]


# ---------- 标题 / 纯文本提取（供 build_kb 登记文档时使用） ----------

def extract_title(text: str, fmt: str) -> str | None:
    """从文档结构提取标题：md/txt 取首个 `#` 标题，html 取首个 h1（回退 <title>）。

    Args:
        text: 文档文本。
        fmt: 格式名（md/txt/html）。

    Returns:
        标题文本；无法识别时 None。
    """
    fmt = (fmt or "").lower()
    if fmt in ("md", "txt"):
        sections = parse_md_sections(text)
        if sections and sections[0]["heading_path"]:
            return sections[0]["heading_path"].splitlines()[0].lstrip("#").strip()
        for line in text.splitlines():
            if line.strip():
                return line.strip()
        return None
    if fmt == "html":
        from bs4 import BeautifulSoup  # 惰性导入

        soup = BeautifulSoup(text, "html.parser")
        h1 = soup.find("h1")
        if h1 and h1.get_text(strip=True):
            return h1.get_text(strip=True)
        title = soup.find("title")
        if title and title.get_text(strip=True):
            return title.get_text(strip=True)
        return None
    return None


def extract_title_file(path: str | Path) -> str | None:
    """从文件提取标题：docx 取首个「标题 1」文本，其余按格式提取。"""
    p = Path(path)
    fmt = _SUFFIX_TO_FORMAT.get(p.suffix.lower())
    if fmt is None:
        return None
    if fmt == "docx":
        sections = parse_docx_sections(p)
        if sections and sections[0]["heading_path"]:
            return sections[0]["heading_path"].splitlines()[0].lstrip("#").strip()
        return None
    return extract_title(p.read_text(encoding="utf-8"), fmt)


def extract_plain_text(path: str | Path) -> str:
    """读取文件纯文本：docx 提取段落文本，其余按 utf-8 读原文（供字符切分策略用）。"""
    p = Path(path)
    if p.suffix.lower() == ".docx":
        doc = _open_docx(p)
        return "\n".join(para.text for para in doc.paragraphs if para.text.strip())
    return p.read_text(encoding="utf-8")
